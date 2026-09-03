"""
DOCX parser -- extracts clause text from a Word document by paragraph
range, per the sidecar manifest's locator ({"para_start": N, "para_end": N},
0-indexed inclusive, matching python-docx's own paragraph indexing).

Reads document.paragraphs -- ordinary body paragraphs only. Does not
handle text inside tables (that's the spreadsheet/table path, not a Word
document's job in this corpus), headers/footers, or tracked-changes
markup (python-docx exposes the current/accepted text, not the revision
history) -- a real Word doc can carry all of those, and this parser
covers what a plain policy-manual-style document actually uses.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from ingestion.formats.manifest import load_manifest
from ingestion.logging_setup import get_logger
from ingestion.schema import Chunk, ChunkMetadata

_log = get_logger("ingestion.formats.docx_parser")


def parse_docx(docx_path: Path, manifest_path: Path, *, repo_root: Path | None = None) -> list[Chunk]:
    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs
    entries = load_manifest(manifest_path)
    source_file = str(docx_path.relative_to(repo_root)) if repo_root else str(docx_path)

    chunks: list[Chunk] = []
    for order, entry in enumerate(entries):
        locator = entry["locator"]
        para_start = locator["para_start"]
        para_end = locator["para_end"]
        if para_start < 0 or para_end >= len(paragraphs) or para_start > para_end:
            raise ValueError(
                f"{manifest_path}: clause {entry['clause_id']!r} locator "
                f"para_start={para_start} para_end={para_end} out of range "
                f"for a {len(paragraphs)}-paragraph document"
            )
        texts = [paragraphs[i].text.strip() for i in range(para_start, para_end + 1)]
        body = "\n".join(t for t in texts if t).strip()
        if not body:
            _log.error("%s: no text in paragraphs %d-%d", entry["clause_id"], para_start, para_end)
            raise ValueError(f"{entry['clause_id']}: no text in paragraphs {para_start}-{para_end}")

        metadata_fields = {k: v for k, v in entry.items() if k not in ("locator",)}
        metadata = ChunkMetadata.model_validate(metadata_fields)
        chunks.append(Chunk(metadata=metadata, body=body, source_file=source_file, order_in_file=order))

    _log.info("parsed %d clauses from DOCX %s", len(chunks), docx_path)
    return chunks
