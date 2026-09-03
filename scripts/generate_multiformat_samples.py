#!/usr/bin/env python3
"""
Generates the real (non-markdown) sample documents used to prove the
format parsers in ingestion/formats/ against actual PDF/DOCX/XLSX/CSV
files, not synthetic text pretending to be one.

These are a SMALL representative slice, not a full corpus rewrite -- see
the conversation with Prashanth on 3 Sep 2026 (README/PROJECT_PLAN.md
record the decision): rewriting all 72 clauses into every format would be
the same text in a different wrapper, not a stronger engineering claim.
What proves something is real parsers tested against real files.

Text content below is copied verbatim from the existing, already-verified
markdown corpus (see corpus/tier1_law/india/india_law.md,
corpus/tier2_policy/india/meridian_india_policy.md,
corpus/tier2_policy/uae/meridian_uae_policy.md) -- nothing new is being
asserted as fact here, only re-packaged into a different file format.

Run once: .venv/bin/python scripts/generate_multiformat_samples.py
Output is committed to git (small files, and they need to exist for the
parser tests to run) under corpus_samples/multi_format/ -- deliberately
NOT under corpus/, so coverage_audit.py and the scored probe set never see
these FMT-* sample clause ids; they are a parser proof, not part of the
graded corpus.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from fpdf import FPDF
from openpyxl import Workbook

REPO_ROOT = Path(__file__).resolve().parent.parent
OUT = REPO_ROOT / "corpus_samples" / "multi_format"


def build_pdf():
    clauses = [
        (
            "FMT-PDF-IN-GRAT-S4-ELIG",
            "Payment of Gratuity Act, 1972 -- Section 4(1)",
            "Gratuity shall be payable to an employee on the termination of his employment "
            "after he has rendered continuous service for not less than five years -- (a) on "
            "his superannuation, or (b) on his retirement or resignation, or (c) on his death "
            "or disablement due to accident or disease. Provided that the completion of "
            "continuous service of five years shall not be necessary where the termination of "
            "the employment of any employee is due to death or disablement.",
        ),
        (
            "FMT-PDF-IN-GRAT-S4-FORMULA",
            "Payment of Gratuity Act, 1972 -- Section 4(2)",
            "For every completed year of service or part thereof in excess of six months, the "
            "employer shall pay gratuity to an employee at the rate of fifteen days' wages "
            "(based on the rate of wages last drawn, comprising basic wage and dearness "
            "allowance) for every completed year of service, calculated as: (last drawn "
            "monthly wage / 26) x 15 x number of completed years of service.",
        ),
        (
            "FMT-PDF-IN-GRAT-S4-CEILING",
            "Payment of Gratuity Act, 1972 (as amended 2018) -- Section 4(3)",
            "The maximum amount of gratuity payable under the Act shall not exceed Rs. "
            "20,00,000 (twenty lakh rupees), as notified by the Central Government with "
            "effect from 29 March 2018, revising the earlier statutory ceiling of Rs. "
            "10,00,000 (ten lakh rupees) fixed in 2010.",
        ),
    ]

    pdf = FPDF()
    pdf.set_font("Helvetica", size=11)
    for clause_id, section, body in clauses:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 13)
        pdf.multi_cell(0, 8, section)
        pdf.ln(2)
        pdf.set_font("Helvetica", size=11)
        pdf.multi_cell(0, 7, body)

    out_dir = OUT / "pdf"
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / "india_gratuity_law.pdf"
    pdf.output(str(pdf_path))

    manifest = [
        {
            "clause_id": cid,
            "locator": {"page_start": i + 1, "page_end": i + 1},
            "country": "India",
            "doc_type": "law",
            "effective_date": "2018-03-29" if "CEILING" in cid else "1972-09-16",
            "temporal_applicability": "POINT_IN_TIME",
            "normative": True,
            "lineage_id": cid,
            "jurisdiction_scope": "india-national",
            "source_url": "corpus/tier1_law/india/india_law.md (same text, PDF format-parser sample)",
        }
        for i, (cid, _, _) in enumerate(clauses)
    ]
    (out_dir / "india_gratuity_law.manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"wrote {pdf_path} + manifest ({len(clauses)} clauses)")


def build_docx():
    clauses = [
        (
            "FMT-DOCX-MER-IN-NOTICE-SENIOR",
            "4.1 Notice period -- Grade M1 and above",
            "An employee in Grade M1 or above who wishes to resign shall give the Company "
            "forty-five (45) days' written notice, or payment of basic salary in lieu of the "
            "unexpired portion of that notice. The Company shall give the same period of "
            "notice where it terminates the engagement of such an employee otherwise than "
            "for misconduct.",
        ),
        (
            "FMT-DOCX-MER-IN-NOTICE-JUNIOR",
            "4.2 Notice period -- Grade A1 to A3",
            "An employee in Grade A1 to A3 who wishes to resign shall give the Company "
            "fifteen (15) days' written notice, or payment of basic salary in lieu of the "
            "unexpired portion of that notice. The same period shall apply where the Company "
            "terminates the engagement of such an employee otherwise than for misconduct.",
        ),
    ]

    doc = Document()
    doc.add_heading("Meridian India HR Policy Manual -- Chapter 4 (format-parser sample)", level=1)
    locators = []
    for cid, heading, body in clauses:
        para_start = len(doc.paragraphs)
        doc.add_paragraph(heading, style="Heading 2")
        doc.add_paragraph(body)
        para_end = len(doc.paragraphs) - 1
        locators.append((cid, heading, para_start, para_end))

    out_dir = OUT / "docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "meridian_india_notice.docx"
    doc.save(str(docx_path))

    manifest = [
        {
            "clause_id": cid,
            "locator": {"para_start": para_start, "para_end": para_end},
            "country": "India",
            "doc_type": "policy",
            "source_doc": "Meridian India HR Policy Manual (DOCX format-parser sample)",
            "effective_date": "2023-01-01",
            "temporal_applicability": "POINT_IN_TIME",
            "normative": True,
            "lineage_id": cid,
            "jurisdiction_scope": "india-national",
        }
        for cid, heading, para_start, para_end in locators
    ]
    (out_dir / "meridian_india_notice.manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"wrote {docx_path} + manifest ({len(clauses)} clauses)")


_HOUSING_ROWS = [
    ("A1-A3", "under 3 years", 4000),
    ("A1-A3", "3 to 5 years", 4500),
    ("A1-A3", "over 5 years", 5000),
    ("M1", "under 3 years", 6500),
    ("M1", "3 to 5 years", 7200),
    ("M1", "over 5 years", 8000),
    ("M2", "under 3 years", 8500),
    ("M2", "3 to 5 years", 9500),
    ("M2", "over 5 years", 10500),
    ("M3", "under 3 years", 11000),
    ("M3", "3 to 5 years", 12500),
    ("M3", "over 5 years", 14000),
    ("D1 and above", "any length", 18000),
]


def build_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Schedule 2"
    ws.append(["Grade", "Years of service", "Monthly housing allowance (AED)"])
    for grade, tenure, amount in _HOUSING_ROWS:
        ws.append([grade, tenure, amount])

    out_dir = OUT / "xlsx"
    out_dir.mkdir(parents=True, exist_ok=True)
    xlsx_path = out_dir / "meridian_uae_housing.xlsx"
    wb.save(str(xlsx_path))

    manifest = [
        {
            "clause_id": "FMT-XLSX-MER-AE-HOUSING",
            "title": "Schedule 2 -- Monthly housing allowance, by grade and completed years of service (AED).",
            "locator": {"sheet_name": "Schedule 2", "header_row": 1, "row_start": 2, "row_end": len(_HOUSING_ROWS) + 1},
            "country": "UAE",
            "doc_type": "policy",
            "source_doc": "Meridian UAE HR Policy Manual -- Schedule 2 (XLSX format-parser sample)",
            "effective_date": "2024-01-01",
            "temporal_applicability": "POINT_IN_TIME",
            "normative": True,
            "lineage_id": "FMT-XLSX-MER-AE-HOUSING",
            "jurisdiction_scope": "uae-mainland",
        }
    ]
    (out_dir / "meridian_uae_housing.manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"wrote {xlsx_path} + manifest ({len(_HOUSING_ROWS)} rows)")


def build_csv():
    out_dir = OUT / "csv"
    out_dir.mkdir(parents=True, exist_ok=True)
    csv_path = out_dir / "meridian_uae_housing.csv"
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Grade", "Years of service", "Monthly housing allowance (AED)"])
        for row in _HOUSING_ROWS:
            writer.writerow(row)

    manifest = [
        {
            "clause_id": "FMT-CSV-MER-AE-HOUSING",
            "title": "Schedule 2 -- Monthly housing allowance, by grade and completed years of service (AED).",
            "locator": {"header_row": 1, "row_start": 2, "row_end": len(_HOUSING_ROWS) + 1},
            "country": "UAE",
            "doc_type": "policy",
            "source_doc": "Meridian UAE HR Policy Manual -- Schedule 2 (CSV format-parser sample, same data as the XLSX sample)",
            "effective_date": "2024-01-01",
            "temporal_applicability": "POINT_IN_TIME",
            "normative": True,
            "lineage_id": "FMT-CSV-MER-AE-HOUSING",
            "jurisdiction_scope": "uae-mainland",
        }
    ]
    (out_dir / "meridian_uae_housing.manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"wrote {csv_path} + manifest ({len(_HOUSING_ROWS)} rows)")


if __name__ == "__main__":
    build_pdf()
    build_docx()
    build_xlsx()
    build_csv()
